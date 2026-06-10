"""
Smart Learn Word 文档生成工具
=============================
支持增量更新：用户每学完一步，文档追加一步内容。
依赖: pip install python-docx
"""

import sys
import os
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _style_doc(doc):
    """设置文档默认样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _add_heading(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return heading


def _add_code_block(doc, code_text):
    """添加代码块（灰底等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 灰色背景通过 shading 实现
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0',
    })
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    return p


def _add_table(doc, headers, rows):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)

    doc.add_paragraph()  # 表后空行
    return table


def init_doc(topic, output_dir):
    """创建学习文档，返回文档路径"""
    if not HAS_DOCX:
        print(json.dumps({"status": "no_docx", "msg": "请先 pip install python-docx"}))
        return None

    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"{topic}_学习笔记.docx")

    doc = Document()
    _style_doc(doc)

    # 封面标题
    title = doc.add_heading(f'{topic} — 学习笔记', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f'创建时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}  |  方法：费曼学习法五步闭环')
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()
    doc.save(filepath)

    print(json.dumps({"status": "ok", "action": "init", "path": filepath}, ensure_ascii=False))
    return filepath


def add_step(filepath, step_num, step_title, content_md=None, content_file=None, extra=None):
    """增量追加一个学习步骤。content_md 或 content_file 二选一"""
    if not HAS_DOCX:
        print(json.dumps({"status": "no_docx", "msg": "请先 pip install python-docx"}))
        return

    if not os.path.exists(filepath):
        print(json.dumps({"status": "error", "msg": f"文档不存在: {filepath}"}))
        return

    # 支持从文件读取内容
    if content_file and os.path.exists(content_file):
        content_md = Path(content_file).read_text(encoding='utf-8')
    if not content_md:
        content_md = ""

    doc = Document(filepath)

    # 步骤标题
    _add_heading(doc, f'步骤{step_num}：{step_title}', level=1)

    # 内容：按段落解析
    for line in content_md.strip().split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('### '):
            _add_heading(doc, line[4:], level=3)
        elif line.startswith('## '):
            _add_heading(doc, line[3:], level=2)
        elif line.startswith('```'):
            pass  # 代码块开始/结束标记跳过，内容在下个循环中处理
        elif line.startswith('|') and '|' in line[1:]:
            pass  # 表格在 extra 中专门处理
        elif line.startswith('- **') and '**：' in line:
            # 薄弱点格式: - **名称**：说明
            p = doc.add_paragraph()
            parts = line[2:].split('**：', 1)
            if len(parts) == 2:
                run_b = p.add_run(parts[0] + '**')
                run_b.font.bold = True
                p.add_run('：' + parts[1])
            else:
                p.add_run(line[2:])
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            doc.add_paragraph(line)

    # 如果有表格数据
    if extra and extra.get('table'):
        tbl = extra['table']
        _add_table(doc, tbl.get('headers', []), tbl.get('rows', []))

    # 如果有代码块
    if extra and extra.get('code'):
        _add_code_block(doc, extra['code'])

    doc.save(filepath)
    print(json.dumps({"status": "ok", "action": "add_step", "step": step_num, "path": filepath}, ensure_ascii=False))


def finalize_doc(filepath, summary, weak_points, keywords):
    """添加总结并最终保存"""
    if not HAS_DOCX:
        print(json.dumps({"status": "no_docx", "msg": "请先 pip install python-docx"}))
        return

    if not os.path.exists(filepath):
        print(json.dumps({"status": "error", "msg": f"文档不存在: {filepath}"}))
        return

    doc = Document(filepath)

    _add_heading(doc, '总结', level=1)

    # 薄弱点
    if weak_points:
        _add_heading(doc, '⚠️ 薄弱点', level=2)
        for wp in weak_points.split(';'):
            wp = wp.strip()
            if wp:
                doc.add_paragraph(wp, style='List Bullet')

    # 关键词
    if keywords:
        _add_heading(doc, '🏷️ 关键词', level=2)
        p = doc.add_paragraph()
        for kw in keywords.split(','):
            kw = kw.strip()
            if kw:
                run = p.add_run(kw)
                run.font.highlight_color = None
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xcc)
                p.add_run('  ')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'— 学习完成于 {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(filepath)
    print(json.dumps({"status": "ok", "action": "finalize", "path": filepath}, ensure_ascii=False))


# ═══════════════════════════════════════
#  CLI 入口
# ═══════════════════════════════════════

def main():
    if not HAS_DOCX:
        print(json.dumps({"status": "no_docx", "msg": "请先 pip install python-docx"}))
        sys.exit(0)

    if len(sys.argv) < 2:
        print("用法: docx_utils.py <command> [options]")
        print("  init       --topic X --output-dir Y")
        print("  add-step   --file X --step N --title Y [--content Z | --content-file F] [--extra JSON]")
        print("  finalize   --file X --summary Y --weak-points Z --keywords W")
        sys.exit(1)

    cmd = sys.argv[1]
    args = _parse_args(sys.argv[2:])

    if cmd == "init":
        filepath = init_doc(
            topic=args.get("--topic", "未命名"),
            output_dir=args.get("--output-dir", "."),
        )

    elif cmd == "add-step":
        extra = json.loads(args["--extra"]) if "--extra" in args else None
        add_step(
            filepath=args["--file"],
            step_num=int(args.get("--step", 1)),
            step_title=args.get("--title", ""),
            content_md=args.get("--content", ""),
            content_file=args.get("--content-file", None),
            extra=extra,
        )

    elif cmd == "finalize":
        finalize_doc(
            filepath=args["--file"],
            summary=args.get("--summary", ""),
            weak_points=args.get("--weak-points", ""),
            keywords=args.get("--keywords", ""),
        )

    else:
        print(f"未知命令: {cmd}")
        sys.exit(1)


def _parse_args(argv):
    result = {}
    i = 0
    while i < len(argv):
        if argv[i].startswith("--"):
            key = argv[i]
            if i + 1 < len(argv) and not argv[i + 1].startswith("--"):
                result[key] = argv[i + 1]
                i += 2
            else:
                result[key] = ""
                i += 1
        else:
            i += 1
    return result


if __name__ == "__main__":
    main()
